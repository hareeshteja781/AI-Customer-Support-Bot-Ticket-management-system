import csv
import io
import os
from pathlib import Path
from typing import Annotated

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from sqlalchemy.orm import Session

from app.api.routes.auth import get_current_user
from app.core.config import UPLOAD_MAX_SIZE
from app.core.database import get_db
from app.models.knowledge_document import KnowledgeDocument
from app.models.user import User
from app.services.rag_service import RAGService

router = APIRouter()

ALLOWED_EXTENSIONS = {"txt", "pdf", "docx", "csv", "xlsx"}
UPLOAD_DIR = Path(__file__).resolve().parents[2] / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def normalize_document_record(document: KnowledgeDocument) -> dict:
    return {
        "id": document.id,
        "filename": document.filename,
        "content_type": document.content_type,
        "status": document.status,
        "storage_path": document.storage_path,
        "text_preview": document.text_preview,
        "uploaded_by": document.uploaded_by,
        "created_at": document.created_at.isoformat() if document.created_at else None,
        "updated_at": document.updated_at.isoformat() if document.updated_at else None,
    }


def extract_text(filename: str, content_type: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    if ext == "csv":
        rows = csv.reader(io.StringIO(file_bytes.decode("utf-8", errors="ignore")))
        return "\n".join(" | ".join(row) for row in rows)
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Unable to read PDF: {exc}") from exc
    if ext == "docx":
        try:
            from docx import Document
            document = Document(io.BytesIO(file_bytes))
            return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Unable to read DOCX: {exc}") from exc
    if ext == "xlsx":
        try:
            from openpyxl import load_workbook
            workbook = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            parts = []
            for sheet in workbook.worksheets:
                parts.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        parts.append(" | ".join(values))
            return "\n".join(parts)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Unable to read XLSX: {exc}") from exc
    raise HTTPException(status_code=400, detail="unsupported file type")


@router.get("")
def list_documents(
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
):
    documents = (
        db.query(KnowledgeDocument)
        .filter(KnowledgeDocument.uploaded_by == current_user.id)
        .order_by(KnowledgeDocument.created_at.desc())
        .all()
    )
    return [normalize_document_record(document) for document in documents]


@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    current_user: Annotated[User, Depends(get_current_user)] = None,
    db: Annotated[Session, Depends(get_db)] = None,
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="missing filename")

    ext = file.filename.rsplit(".", 1)[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="unsupported file type")

    file_bytes = await file.read()
    if len(file_bytes) > UPLOAD_MAX_SIZE:
        raise HTTPException(status_code=400, detail="file too large")

    text = extract_text(file.filename, file.content_type or "", file_bytes).strip()
    if not text:
        raise HTTPException(status_code=400, detail="no readable text was found in the document")

    safe_name = Path(file.filename).name.replace("/", "_").replace("\\", "_")
    storage_path = UPLOAD_DIR / safe_name
    storage_path.write_bytes(file_bytes)

    chunks = RAGService().chunk_text(text)
    preview = chunks[0][:1000] if chunks else text[:1000]

    document = KnowledgeDocument(
        filename=file.filename,
        content_type=file.content_type or "application/octet-stream",
        status="ready",
        uploaded_by=current_user.id,
        storage_path=str(storage_path),
        text_preview="\n\n".join(chunks[:5]),
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    return {
        **normalize_document_record(document),
        "type": document.content_type,
        "preview": preview,
    }


@router.delete("/{document_id}")
def delete_document(
    document_id: int,
    current_user: Annotated[User, Depends(get_current_user)],
    db: Annotated[Session, Depends(get_db)],
):
    document = (
        db.query(KnowledgeDocument)
        .filter(KnowledgeDocument.id == document_id, KnowledgeDocument.uploaded_by == current_user.id)
        .first()
    )
    if not document:
        raise HTTPException(status_code=404, detail="document not found")
    if document.storage_path and os.path.exists(document.storage_path):
        os.remove(document.storage_path)
    db.delete(document)
    db.commit()
    return {"status": "deleted", "document_id": document_id}
